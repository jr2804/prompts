# /// script
# requires-python = ">=3.11"
# dependencies = [
#   "python-docx>=1.1.2",
#   "docxtpl>=0.19.0",
# ]
# ///
"""SVG image support for python-docx and docxtpl.

This module monkey-patches python-docx to add SVG image support and provides
a custom ``SvgInlineImage`` class for use with docxtpl templates.

The monkey-patch is applied **automatically** when this module is imported ---
no need to worry about import order. Just import ``SvgInlineImage`` (or any
symbol from this module) and the patch will be active for all subsequent
python-docx / docxtpl operations.

Based on:
- https://gist.github.com/Kladdy/d3bdb9bbf2c38d4f194ea9a7904fc3f2
- https://github.com/python-openxml/python-docx/pull/1107

Usage
-----

**Direct python-docx usage** (add SVG inline shape to a paragraph):

    import docx
    import svg  # noqa: F401  -- patch is active on import

    doc = docx.Document()
    doc.add_picture("chart.svg", width=docx.shared.Mm(100))
    doc.save("output.docx")

**docxtpl template usage** (pass SVG to a Jinja2 template):

    from docxtpl import DocxTemplate
    from scripts.svg import SvgInlineImage
    from docx.shared import Mm

    tpl = DocxTemplate("template.docx")
    img = SvgInlineImage(tpl, "chart.svg", width=Mm(100))
    tpl.render({"chart": img})
    tpl.save("output.docx")

**Check if file is SVG:**

    from scripts.svg import is_svg
    assert is_svg("icon.svg")
"""

from __future__ import annotations

import xml.etree.ElementTree as ET
from io import BytesIO
from pathlib import Path

import docx
from docx.image.constants import MIME_TYPE
from docx.image.exceptions import UnrecognizedImageError
from docx.image.image import BaseImageHeader
from docx.shared import Mm
from docxtpl import DocxTemplate, InlineImage


# ---------------------------------------------------------------------------
# Patch: extended image header factory that recognises SVG
# ---------------------------------------------------------------------------

def _ImageHeaderFactory(stream: BytesIO) -> BaseImageHeader:
    """Return a |BaseImageHeader| subclass for the image in *stream*.

    This patched version adds SVG support to python-docx's standard image
    header detection.  It reads the first 256 bytes (instead of 64) so that
    XML-declared SVGs are recognised.

    Raises:
        UnrecognizedImageError: If the image format cannot be determined.
    """
    from docx.image import SIGNATURES

    def _read_256(s: BytesIO) -> bytes:
        s.seek(0)
        return s.read(256)

    header = _read_256(stream)
    for cls, offset, signature_bytes in SIGNATURES:
        end = offset + len(signature_bytes)
        if header[offset:end] == signature_bytes:
            return cls.from_stream(stream)
    raise UnrecognizedImageError


# ---------------------------------------------------------------------------
# SVG image header parser
# ---------------------------------------------------------------------------

class Svg(BaseImageHeader):
    """Image header parser for SVG images.

    Extracts pixel dimensions from the SVG's ``width`` / ``height``
    attributes or, when those are absent, from the ``viewBox``.
    """

    @classmethod
    def from_stream(cls, stream: BytesIO) -> Svg:
        """Construct |Svg| from an SVG byte stream."""
        px_width, px_height = cls._dimensions_from_stream(stream)
        return cls(px_width, px_height, 72, 72)

    @property
    def content_type(self) -> str:
        return MIME_TYPE.SVG  # type: ignore[attr-defined]

    @property
    def default_ext(self) -> str:
        return "svg"

    @classmethod
    def _dimensions_from_stream(cls, stream: BytesIO) -> tuple[int, int]:
        """Extract (width, height) in pixels from an SVG stream.

        Handles:
        * Float dimension values (e.g. matplotlib ``"457.774606pt"``)
        * Missing ``width`` / ``height`` -> fallback to ``viewBox``
        * Units: pt, px, cm, mm, %, em, in
        """
        stream.seek(0)
        data = stream.read()
        root = ET.fromstring(data)  # noqa: S314

        # -- explicit width / height attributes --
        if "width" in root.attrib and "height" in root.attrib:
            width = _parse_dimension(root.attrib["width"])
            height = _parse_dimension(root.attrib["height"])

        # -- viewBox fallback --
        elif "viewBox" in root.attrib:
            parts = root.attrib["viewBox"].split()
            if len(parts) >= 4:
                width = round(float(parts[2]))
                height = round(float(parts[3]))
            else:
                width, height = 400, 300
        else:
            width, height = 400, 300

        return width, height


def _parse_dimension(value: str) -> int:
    """Strip units from an SVG dimension string and return pixels as int."""
    import re

    num = re.sub(r"[^\d.]", "", value)
    try:
        return round(float(num))
    except ValueError:
        return 400


# ==========================================================================
# MONKEY-PATCH APPLICATION (applied on import)
# ==========================================================================

docx.image.Svg = Svg  # type: ignore[attr-defined]
docx.image.constants.MIME_TYPE.SVG = "image/svg+xml"  # type: ignore[attr-defined]

_SVG_SIGNATURES = [
    (Svg, 0, b"<?xml version="),   # XML declaration first
    (Svg, 0, b"<svg "),            # bare SVG start
    (Svg, 0, b"<!DOCTYPE svg"),    # DOCTYPE preamble
]
docx.image.SIGNATURES = tuple(list(docx.image.SIGNATURES) + _SVG_SIGNATURES)  # type: ignore[attr-defined]
docx.image.image._ImageHeaderFactory = _ImageHeaderFactory  # type: ignore[attr-defined]


# ==========================================================================
# SvgInlineImage -- docxtpl-compatible InlineImage with SVG support
# ==========================================================================

class SvgInlineImage(InlineImage):
    """Extended ``InlineImage`` for **docxtpl** that accepts SVG files.

    Automatically detects SVG vs. bitmap images.  When both *width* and
    *height* are omitted the image is sized to the usable page width
    (page width minus left and right margins).

    Example::

        from docxtpl import DocxTemplate
        from docx.shared import Mm
        from scripts.svg import SvgInlineImage

        tpl = DocxTemplate("template.docx")
        img = SvgInlineImage(tpl, "chart.svg", width=Mm(100))
        tpl.render({"chart": img})
        tpl.save("output.docx")
    """

    def __init__(
        self,
        tpl: DocxTemplate,
        image_descriptor: Path | str | BytesIO,
        width: float | int | None = None,
        height: float | int | None = None,
    ) -> None:
        if width is None and height is None:
            width = self._get_page_width(tpl)

        # Convert Path -> BytesIO for SVGs so the patch can parse them.
        if isinstance(image_descriptor, Path):
            if image_descriptor.suffix.lower() == ".svg":
                image_descriptor = BytesIO(image_descriptor.read_bytes())
            else:
                image_descriptor = str(image_descriptor)

        super().__init__(tpl, image_descriptor, width, height)

    @staticmethod
    def _get_page_width(tpl: DocxTemplate) -> Mm:
        """Usable page width = page width - left margin - right margin."""
        doc = docx.Document(str(tpl.template_file))
        section = doc.sections[0]

        pw = section.page_width or 0
        lm = section.left_margin or 0
        rm = section.right_margin or 0

        usable_emu = pw - lm - rm
        return Mm(usable_emu / 36000)


# ==========================================================================
# Utilities
# ==========================================================================

def is_svg(file_path: Path) -> bool:
    """Return ``True`` if *file_path* looks like an SVG image.

    Checks both file extension (``.svg``) and content signature.
    """
    if not file_path.exists():
        return False
    if file_path.suffix.lower() == ".svg":
        return True
    try:
        header = file_path.read_bytes()[:256]
        return any(tag in header for tag in (b"<svg ", b"<?xml ", b"<!DOCTYPE svg"))
    except OSError:
        return False


# ==========================================================================
# Run as script: apply patch + dump diagnostic info
# ==========================================================================
if __name__ == "__main__":
    import sys

    print("SVG patch applied to python-docx.")
    print(f"  MIME_TYPE.SVG = {MIME_TYPE.SVG}")
    print(f"  SVG signatures registered: {len(_SVG_SIGNATURES)}")
    print(f"  SvgInlineImage available: {SvgInlineImage is not None}")
    print(f"  is_svg helper available: {is_svg is not None}")

    if len(sys.argv) > 1:
        for arg in sys.argv[1:]:
            p = Path(arg)
            if p.exists() and is_svg(p):
                stream = BytesIO(p.read_bytes())
                dims = Svg._dimensions_from_stream(stream)
                print(f"  {p.name}: {dims[0]}x{dims[1]} px")
            else:
                print(f"  {p.name}: not an SVG file")
